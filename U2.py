import os
import re
import sys
import fitz
from docx import Document
from docx.shared import Pt
from dotenv import load_dotenv
from google import genai
from google.genai import types
from google.cloud import bigquery
from google.oauth2 import service_account

# ==============================
# LOAD ENV VARIABLES
# ==============================
load_dotenv()
GEMINI_API_KEY = os.getenv("GEMINI_API_KEY")
if not GEMINI_API_KEY:
    raise ValueError("❌ GEMINI_API_KEY not found.")
client = genai.Client(api_key=GEMINI_API_KEY)

BQ_SERVICE_ACCOUNT = os.getenv("BQ_SERVICE_ACCOUNT")
BQ_PROJECT_ID      = os.getenv("BQ_PROJECT_ID")
BQ_DATASET_ID      = os.getenv("BQ_DATASET_ID")
BQ_BRANDS_TABLE    = os.getenv("BQ_BRANDS_TABLE")

for var_name, var_val in [("BQ_SERVICE_ACCOUNT", BQ_SERVICE_ACCOUNT),
                          ("BQ_PROJECT_ID", BQ_PROJECT_ID),
                          ("BQ_DATASET_ID", BQ_DATASET_ID),
                          ("BQ_BRANDS_TABLE", BQ_BRANDS_TABLE)]:
    if not var_val:
        raise ValueError(f"❌ {var_name} not found in .env")

bq_credentials = service_account.Credentials.from_service_account_file(
    BQ_SERVICE_ACCOUNT,
    scopes=["https://www.googleapis.com/auth/bigquery"]
)
bq_client = bigquery.Client(project=BQ_PROJECT_ID, credentials=bq_credentials)

# ==============================
# PARSE CLI ARGUMENTS
# ==============================
def parse_drugs():
    """
    Reads drug names from CLI args.
    Usage: python code.py DrugA DrugB DrugC ...
    Falls back to a prompt if no args supplied.
    """
    if len(sys.argv) > 1:
        drugs = sys.argv[1:]
    else:
        user_input = input(
            "No drugs supplied. Enter drug names separated by spaces: "
        ).strip()
        if not user_input:
            raise SystemExit("❌ No drug names provided. Exiting.")
        drugs = user_input.split()
    print(f"💊 Drugs to analyse: {', '.join(drugs)}")
    return drugs

# ==============================
# LOOKUP MOA FROM BIGQUERY
# ==============================
def lookup_moa(drug_names: list[str]) -> dict[str, dict]:
    """
    Queries BigQuery to map each drug name to its MoA.
    Returns dict: {drug_name: {"moa": ..., "moa_detailed": ...}}
    Falls back to "Unknown" if a drug is not found.
    """
    table_ref = f"`{BQ_PROJECT_ID}.{BQ_DATASET_ID}.{BQ_BRANDS_TABLE}`"

    # Build parameterised query
    placeholders = ", ".join(f"@drug_{i}" for i in range(len(drug_names)))
    query = f"""
        SELECT
            Cleaned_Generic_Name,
            Mechanism_of_Action,
            Mechanism_of_Action_Detailed
        FROM {table_ref}
        WHERE LOWER(Cleaned_Generic_Name) IN ({placeholders})
    """
    job_config = bigquery.QueryJobConfig(
        query_parameters=[
            bigquery.ScalarQueryParameter(f"drug_{i}", "STRING", name.lower())
            for i, name in enumerate(drug_names)
        ]
    )

    results = bq_client.query(query, job_config=job_config).result()

    # Build lookup keyed by lower-case generic name
    moa_map_lower: dict[str, dict] = {}
    for row in results:
        moa_map_lower[row.Cleaned_Generic_Name.lower()] = {
            "moa":          row.Mechanism_of_Action or "Unknown",
            "moa_detailed": row.Mechanism_of_Action_Detailed or "",
        }

    # Map back to the original drug names supplied by the user
    drug_moa: dict[str, dict] = {}
    for drug in drug_names:
        match = moa_map_lower.get(drug.lower())
        if match:
            drug_moa[drug] = match
        else:
            print(f"⚠️  '{drug}' not found in BigQuery table. Using 'Unknown' as MOA.")
            drug_moa[drug] = {"moa": "Unknown", "moa_detailed": ""}

    return drug_moa

# ==============================
# CONFIG
# ==============================
INPUT_PDF  = "soc.pdf"
MODEL_NAME = "gemini-2.5-flash"

# ==============================
# SYSTEM INSTRUCTION
# ==============================
SYSTEM_INSTRUCTION = """You are a Clinical Pharmacist and Diabetes/Obesity Treatment Pathway Analyst.
Analyze the provided Standard of Care (SoC) document and create a country-level Lines of Therapy (LOT) benchmark.

PURPOSE:
Neutral, evidence-based extraction for downstream LOT comparison.

SCOPE:
- Type 2 Diabetes Mellitus (T2DM)
- Obesity / Weight Management

OUTPUT FORMAT STRICT:
Country: [Country]
Country-Level SoC LOT Benchmark

1L:
Recommended therapies/interventions:
Patient segment or trigger:
Treatment goal or rationale:

2L:
Recommended therapies/interventions:
Patient segment or trigger:
Treatment goal or rationale:

3L:
Recommended therapies/interventions:
Patient segment or trigger:
Treatment goal or rationale:

Salvage:
Recommended therapies/interventions:
Patient segment or trigger:
Treatment goal or rationale:

Therapy Classes Explicitly Mentioned:
- List all major classes

FORMATTING RULES:
- Output plain business English.
- Do NOT use markdown.
- Do NOT use LaTeX.
- Use Unicode symbols directly: ≥ ≤ kg/m2.
- Do NOT use *, **, #, or $.
- Produce Word-document-ready text."""

# ==============================
# BUILD OVERLAY PROMPT DYNAMICALLY
# ==============================
def build_overlay_prompt(drugs: list[str], drug_moa: dict[str, dict]) -> str:
    """
    Builds the full overlay prompt using the supplied drug names and their
    MoA looked up from BigQuery.  drug_moa maps each drug name to a dict
    with keys 'moa' and 'moa_detailed'.
    """
    # Numbered molecule list for the task section
    molecule_lines = []
    for i, drug in enumerate(drugs):
        moa = drug_moa[drug]["moa"]
        moa_det = drug_moa[drug]["moa_detailed"]
        detail = f" ({moa_det})" if moa_det else ""
        molecule_lines.append(f"{i+1}. {drug} - MOA: {moa}{detail}")
    molecule_list = "\n".join(molecule_lines)

    # Per-molecule output blocks
    output_blocks = "\n\n".join(
        f"{drug}:\nMolecule: {drug}\nMOA: {drug_moa[drug]['moa']}\nFinal LoT Category:\nRationale:"
        for drug in drugs
    )

    return f"""You are a Senior Market Access Analyst.

Task:
Determine the most appropriate Line of Treatment (LOT) classification for each molecule below based ONLY on the supplied MOA and the SoC definitions provided above.

Molecules and their MOA:
{molecule_list}

LOT ASSIGNMENT RULES:
- Identify the pharmacologic class, therapeutic modality, or treatment approach from the MOA.
- Match the MOA to the therapy classes, mechanisms, modalities, or treatment approaches explicitly described in the SoC.
- Determine the earliest treatment line in which the matching therapy class, mechanism, modality, or treatment approach appears.
- If the MOA aligns with multiple treatment lines, assign the earliest applicable line.
- Do NOT use historical treatment sequencing, current prescribing patterns, external guidelines, prior knowledge, or assumptions outside the supplied SoC.
- Evaluate the MOA against all SoC pathways and patient segments described in the benchmark.
- When a therapy class, mechanism, modality, or treatment approach appears in multiple SoC pathways, prioritize the earliest treatment line in which it is explicitly recommended.
- Only assign Second-Line when the MOA aligns primarily with therapies described as add-on, substitute, escalation, replacement, or post-failure options relative to First-Line treatment.
- Only assign Third-Line when the MOA aligns primarily with therapies described as later-line escalation, refractory-disease management, restricted-use therapies, or options used after failure of earlier treatment lines.
- Only assign Salvage when the MOA aligns primarily with rescue therapies, transplantation, last-resort interventions, or therapies explicitly described in the Salvage section of the SoC.
- When uncertainty exists, assign the earliest treatment line supported by the MOA-to-SoC mapping.

CLASSIFICATION OPTIONS:
- First-line standard of care
- Strong first-line alternative / dominant second-line
- Second-line option
- Third-line or restricted niche use
- Salvage / last-resort use

OUTPUT FORMAT:
{output_blocks}

FORMATTING RULES:
- Output plain business English.
- Do NOT use markdown.
- Do NOT use LaTeX.
- Use Unicode symbols directly.
- Produce Word-document-ready text."""

# ==============================
# PDF EXTRACTION
# ==============================
def extract_pdf_text(pdf_path):
    doc = fitz.open(pdf_path)
    text = ""
    for page in doc:
        text += page.get_text("text")
    doc.close()
    return text

# ==============================
# CLEAN GEMINI OUTPUT
# ==============================
def clean_text(text):
    replacements = {
        "$\\ge$": "≥",
        "$\\le$": "≤",
        "\\ge":   "≥",
        "\\le":   "≤",
        "kg/m2":  "kg/m2",
        "**": "",
        "*":  "",
        "#":  ""
    }
    for old, new in replacements.items():
        text = text.replace(old, new)
    text = re.sub(r"\$+", "", text)
    return text.strip()

# ==============================
# GEMINI CALLS
# ==============================
def run_soc_extraction(pdf_text):
    response = client.models.generate_content(
        model=MODEL_NAME,
        contents=f"Below is the full SoC document:\n\n{pdf_text}",
        config=types.GenerateContentConfig(
            system_instruction=SYSTEM_INSTRUCTION,
            temperature=0.1
        )
    )
    return response.text if response.text else "No response received."

def run_overlay_analysis(soc_text, overlay_prompt):
    response = client.models.generate_content(
        model=MODEL_NAME,
        contents=(
            "Below is the extracted SoC benchmark. "
            "Use it as the sole reference for LOT assignment.\n\n"
            f"{soc_text}\n\n{overlay_prompt}"
        ),
        config=types.GenerateContentConfig(temperature=0.1)
    )
    return response.text if response.text else "No response received."

# ==============================
# WRITE NICELY FORMATTED DOCX
# ==============================
def add_formatted_section(doc, text, drug_names: list[str]):
    """
    Renders extracted text into the Word document.
    drug_names is used to detect per-drug heading lines dynamically.
    """
    drug_headings = {f"{d}:".lower() for d in drug_names}

    for line in text.split("\n"):
        line = clean_text(line)
        if not line:
            continue
        line_lower = line.lower()

        if line.startswith("Country:"):
            doc.add_heading(line, level=1)

        elif line in ["1L:", "2L:", "3L:", "Salvage:"]:
            doc.add_heading(line.replace(":", ""), level=2)

        elif (
            line_lower.startswith("recommended therapies")
            or line_lower.startswith("patient segment")
            or line_lower.startswith("treatment goal")
            or line_lower.startswith("molecule:")
            or line_lower.startswith("moa:")
            or line_lower.startswith("final lot category:")
            or line_lower.startswith("rationale:")
        ):
            p = doc.add_paragraph()
            run = p.add_run(line)
            run.bold = True

        elif line.startswith("-"):
            doc.add_paragraph(line[1:].strip(), style="List Bullet")

        elif line_lower in drug_headings:
            # Dynamic per-drug section heading
            doc.add_heading(line.replace(":", ""), level=2)

        else:
            doc.add_paragraph(line)

# ==============================
# MAIN
# ==============================
def main():
    # 1. Read drug names from CLI (or prompt)
    drugs = parse_drugs()

    # 2. Build dynamic output filename  e.g. soc_lot_DrugA_DrugB.docx
    slug = "_".join(d.lower() for d in drugs)
    output_doc = f"soc_lot_{slug}.docx"

    # 3. Look up MoA from BigQuery
    print("🔍 Looking up MoA from BigQuery...")
    drug_moa = lookup_moa(drugs)
    for drug, info in drug_moa.items():
        detail = f" ({info['moa_detailed']})" if info["moa_detailed"] else ""
        print(f"   {drug} → {info['moa']}{detail}")

    # 4. Build overlay prompt
    overlay_prompt = build_overlay_prompt(drugs, drug_moa)

    # 5. Extract PDF
    print("📄 Reading PDF...")
    pdf_text = extract_pdf_text(INPUT_PDF)

    # 6. SoC extraction
    print("🧠 Extracting SoC...")
    soc_output = run_soc_extraction(pdf_text)

    # 7. Overlay analysis
    print("⚡ Running overlay analysis...")
    overlay_output = run_overlay_analysis(soc_output, overlay_prompt)

    # 8. Write DOCX
    print("📝 Generating DOCX...")
    doc = Document()

    doc.add_heading("Country-Level SoC LOT Benchmark", level=1)
    add_formatted_section(doc, soc_output, drugs)

    doc.add_page_break()

    heading = " & ".join(drugs) + " LOT Classification"
    doc.add_heading(heading, level=1)
    add_formatted_section(doc, overlay_output, drugs)

    doc.save(output_doc)
    print(f"✅ Saved: {output_doc}")

if __name__ == "__main__":
    main()
